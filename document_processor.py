"""
Document Processor Module
=========================
This module handles loading and chunking documents for the RAG system.

Key concepts:
- CHUNKING: We split documents into smaller pieces because:
  1. Embeddings work better on focused text
  2. We can retrieve just the relevant parts, not entire documents
  3. LLMs have context limits, so we need to be selective

- OVERLAP: Chunks overlap slightly so we don't lose context at boundaries.
  Example: If a sentence spans two chunks, overlap ensures we capture it.
"""

import os
import asyncio
from pathlib import Path
from typing import List, Dict
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor


@dataclass
class DocumentChunk:
    """Represents a piece of a document with metadata."""
    content: str           # The actual text
    source: str           # Original filename
    chunk_index: int      # Which chunk this is (0, 1, 2, ...)
    
    def to_dict(self) -> Dict:
        return {
            "content": self.content,
            "source": self.source,
            "chunk_index": self.chunk_index
        }


class DocumentProcessor:
    """
    Loads documents and splits them into chunks for embedding.
    
    Supports: .txt, .md, .pdf, .docx, .xlsx, .csv, .png, .jpg, .jpeg files
    """
    
    # File size limits
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB per file
    MAX_BATCH_SIZE = 50 * 1024 * 1024  # 50MB total per batch
    MAX_FILES_PER_BATCH = 10
    
    def __init__(self, chunk_size: int = 350, chunk_overlap: int = 35):
        """
        Args:
            chunk_size: Target number of characters per chunk (smaller = faster embedding)
            chunk_overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        # Thread pool for parallel file processing
        self._executor = ThreadPoolExecutor(max_workers=4)
    
    def load_document(self, file_path: str) -> str:
        """
        Load a document and return its text content.
        
        Automatically detects file type and uses appropriate loader.
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == ".txt":
            return self._load_text(file_path)
        elif extension == ".md":
            return self._load_text(file_path)  # Markdown is just text
        elif extension == ".pdf":
            return self._load_pdf(file_path)
        elif extension == ".docx":
            return self._load_docx(file_path)
        elif extension == ".xlsx":
            return self._load_xlsx(file_path)
        elif extension == ".csv":
            return self._load_csv(file_path)
        elif extension in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"]:
            return self._load_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _load_text(self, file_path: str) -> str:
        """Load plain text or markdown files."""
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    
    def _load_pdf(self, file_path: str) -> str:
        """
        Load PDF files using pypdf.
        Falls back to Vision AI OCR for scanned PDFs.
        """
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text)
        
        combined_text = "\n\n".join(text_parts)
        
        # If no text extracted, try OCR for scanned PDFs
        if not combined_text.strip():
            print(f"  No text in PDF, attempting Vision AI OCR...")
            try:
                from ocr_service import ocr_service
                ocr_text = ocr_service.extract_text_from_scanned_pdf(file_path, max_pages=10)
                if ocr_text and not ocr_text.startswith("["):
                    return f"[Scanned PDF - text extracted via Vision AI]\n\n{ocr_text}"
            except Exception as e:
                print(f"  Vision AI OCR failed for PDF: {e}")
            
            return f"[PDF with {len(reader.pages)} pages - no extractable text]"
        
        return combined_text
    
    def _load_docx(self, file_path: str) -> str:
        """
        Load Word documents using python-docx.
        """
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    
    def _load_xlsx(self, file_path: str) -> str:
        """
        Load Excel files using openpyxl.
        """
        from openpyxl import load_workbook
        
        wb = load_workbook(file_path, data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in row_values):
                    text_parts.append(" | ".join(row_values))
        
        return "\n".join(text_parts)
    
    def _load_csv(self, file_path: str) -> str:
        """
        Load CSV files.
        """
        import csv
        
        text_parts = []
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                if any(cell.strip() for cell in row):
                    text_parts.append(" | ".join(row))
        
        return "\n".join(text_parts)
    
    def _load_image(self, file_path: str) -> str:
        """
        Load images and extract text using USF Vision API (OCR).
        Falls back to pytesseract if API fails.
        """
        # Try USF Vision API first (better quality)
        try:
            from ocr_service import ocr_service
            
            print(f"  Using USF Vision API for OCR...")
            text = ocr_service.extract_text_from_image(file_path)
            
            if text and not text.startswith("[OCR Error"):
                return f"[Image text extracted via Vision AI]\n\n{text}"
        except Exception as e:
            print(f"  USF Vision API failed: {e}, falling back to pytesseract")
        
        # Fallback to pytesseract
        try:
            import pytesseract
            from PIL import Image
            
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            
            if text.strip():
                return f"[Image text extracted via OCR]\n\n{text}"
            else:
                return f"[Image file: {Path(file_path).name}]\n\nNo text detected in image."
        except Exception as e:
            return f"[Image file: {Path(file_path).name}]\n\nCould not extract text: {str(e)}"
    
    def chunk_text(self, text: str, source: str) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks.
        
        Strategy:
        1. Try to split on paragraph boundaries first
        2. If paragraphs are too long, split on sentences
        3. If sentences are too long, split on words
        
        This preserves semantic meaning better than arbitrary character splits.
        """
        # Clean up the text
        text = text.strip()
        if not text:
            return []
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size
            
            # If we're not at the end, try to find a good break point
            if end < len(text):
                # Look for paragraph break first (best)
                paragraph_break = text.rfind("\n\n", start, end)
                if paragraph_break > start + self.chunk_size // 2:
                    end = paragraph_break
                else:
                    # Look for sentence break (good)
                    for punct in [". ", "! ", "? ", ".\n", "!\n", "?\n"]:
                        sentence_break = text.rfind(punct, start, end)
                        if sentence_break > start + self.chunk_size // 2:
                            end = sentence_break + 1
                            break
                    else:
                        # Look for word break (acceptable)
                        space_break = text.rfind(" ", start, end)
                        if space_break > start + self.chunk_size // 2:
                            end = space_break
            
            # Extract the chunk
            chunk_text = text[start:end].strip()
            
            if chunk_text:  # Only add non-empty chunks
                chunks.append(DocumentChunk(
                    content=chunk_text,
                    source=source,
                    chunk_index=chunk_index
                ))
                chunk_index += 1
            
            # Move start position, accounting for overlap
            start = end - self.chunk_overlap
            if start <= chunks[-1].chunk_index if chunks else 0:
                start = end  # Prevent infinite loop
        
        return chunks
    
    def process_document(self, file_path: str) -> List[DocumentChunk]:
        """
        Full pipeline: load document and split into chunks.
        
        Returns a list of DocumentChunk objects ready for embedding.
        """
        # Get just the filename for the source
        source = Path(file_path).name
        
        # Load the document
        text = self.load_document(file_path)
        
        # Split into chunks
        chunks = self.chunk_text(text, source)
        
        print(f"  Processed '{source}': {len(chunks)} chunks created")
        
        return chunks
    
    async def process_document_async(self, file_path: str) -> List[DocumentChunk]:
        """
        Async version of process_document for parallel processing.
        Runs the CPU-bound work in a thread pool.
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            self._executor,
            self.process_document,
            file_path
        )
    
    @staticmethod
    def validate_file(filename: str, file_size: int) -> Dict:
        """
        Validate a file before upload.
        Returns dict with 'valid' bool and 'error' message if invalid.
        """
        # Check file extension
        allowed_extensions = {'.txt', '.md', '.pdf', '.docx', '.xlsx', '.csv', '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}
        ext = Path(filename).suffix.lower()
        
        if ext not in allowed_extensions:
            return {
                'valid': False,
                'error': f"Unsupported file type: {ext}. Allowed: PDF, DOCX, XLSX, CSV, TXT, MD, PNG, JPG"
            }
        
        # Check file size
        if file_size > DocumentProcessor.MAX_FILE_SIZE:
            size_mb = file_size / (1024 * 1024)
            return {
                'valid': False,
                'error': f"File too large: {size_mb:.1f}MB. Maximum: 10MB per file"
            }
        
        return {'valid': True, 'error': None}
    
    @staticmethod
    def validate_batch(files: List[Dict]) -> Dict:
        """
        Validate a batch of files.
        files: List of {'filename': str, 'size': int}
        Returns dict with 'valid' bool and 'error' message if invalid.
        """
        # Check number of files
        if len(files) > DocumentProcessor.MAX_FILES_PER_BATCH:
            return {
                'valid': False,
                'error': f"Too many files: {len(files)}. Maximum: {DocumentProcessor.MAX_FILES_PER_BATCH} files per upload"
            }
        
        # Check total size
        total_size = sum(f['size'] for f in files)
        if total_size > DocumentProcessor.MAX_BATCH_SIZE:
            total_mb = total_size / (1024 * 1024)
            return {
                'valid': False,
                'error': f"Total size too large: {total_mb:.1f}MB. Maximum: 50MB per batch"
            }
        
        return {'valid': True, 'error': None}


# Quick test if run directly
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test with a sample text
    sample_text = """
    This is the first paragraph. It contains some information about RAG systems.
    RAG stands for Retrieval-Augmented Generation.
    
    This is the second paragraph. It explains how embeddings work.
    Embeddings are numerical representations of text that capture semantic meaning.
    
    This is the third paragraph. It discusses vector databases.
    Vector databases store embeddings and allow fast similarity search.
    """
    
    chunks = processor.chunk_text(sample_text, "test.txt")
    for chunk in chunks:
        print(f"\n--- Chunk {chunk.chunk_index} ---")
        print(chunk.content[:100] + "..." if len(chunk.content) > 100 else chunk.content)
